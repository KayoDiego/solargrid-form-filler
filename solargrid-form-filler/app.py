from flask import Flask, render_template, request, redirect, url_for, session, send_file
from docx import Document
from datetime import datetime
import os

app = Flask(__name__)
app.secret_key = 'chave_secreta_segura'

OUTPUT_FOLDER = 'output'
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

USUARIO = 'admin'
SENHA = '1234'

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        usuario = request.form['usuario']
        senha = request.form['senha']
        if usuario == USUARIO and senha == SENHA:
            session['usuario'] = usuario
            return redirect(url_for('index'))
        else:
            return 'Login inválido!'
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.pop('usuario', None)
    return redirect(url_for('login'))

@app.route('/')
def index():
    if 'usuario' not in session:
        return redirect(url_for('login'))
    return render_template('index.html')

@app.route('/gerar', methods=['POST'])
def gerar():
    if 'usuario' not in session:
        return redirect(url_for('login'))

    # Salva o arquivo Word enviado pelo usuário
    file = request.files['documento']
    file_path = os.path.join(OUTPUT_FOLDER, 'modelo_temp.docx')

    # Se já existir, deleta o antigo
    if os.path.exists(file_path):
        os.remove(file_path)

    file.save(file_path)

    # Abre o Word
    doc = Document(file_path)

    # Gera data atual
    data_hoje = datetime.now().strftime('%d/%m/%Y')

    # Mapeia os campos recebidos
    campos = {
        '{{NOME}}': request.form.get('nome', ''),
        '{{CPF}}': request.form.get('cpf', ''),
        '{{EMAIL}}': request.form.get('email', ''),
        '{{TELEFONE}}': request.form.get('telefone', ''),
        '{{ENDERECO}}': request.form.get('endereco', ''),
        '{{RAZAO_SOCIAL}}': request.form.get('razao_social', ''),
        '{{CNPJ}}': request.form.get('cnpj', ''),
        '{{CEP}}': request.form.get('cep', ''),
        '{{NOME_REPRESENTANTE}}': request.form.get('nome_representante', ''),
        '{{CPF_REPRESENTANTE}}': request.form.get('cpf_representante', ''),
        '{{EMAIL_REPRESENTANTE}}': request.form.get('email_representante', ''),
        '{{TELEFONE_REPRESENTANTE}}': request.form.get('telefone_representante', ''),
        '{{UC}}': request.form.get('uc', ''),
        '{{DESCONTO}}': request.form.get('desconto', ''),
        '{{DATA_ATUAL}}': data_hoje
    }

    # Faz a substituição nos parágrafos
    for p in doc.paragraphs:
        for key, value in campos.items():
            if key in p.text:
                p.text = p.text.replace(key, value)

    # Faz a substituição também dentro de tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in campos.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    output_path = os.path.join(OUTPUT_FOLDER, 'Contrato_Preenchido.docx')
    doc.save(output_path)

    return send_file(output_path, as_attachment=True)

if __name__ == '__main__':
    app.run(debug=True)
